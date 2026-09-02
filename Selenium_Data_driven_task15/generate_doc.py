from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Iris Communication")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)

# Address line
addr = doc.add_paragraph()
addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
addr.add_run("No.582, 4th Main, 4th Cross, Maruthilayout, Vasanthapura, Bengaluru \u2013 560061").font.size = Pt(9)

gst = doc.add_paragraph()
gst.alignment = WD_ALIGN_PARAGRAPH.CENTER
gst.add_run("GST: 29AAIFI7005B1ZN").font.size = Pt(9)

# Heading
heading = doc.add_paragraph()
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
hrun = heading.add_run("Authorisation letter")
hrun.bold = True
hrun.underline = True
hrun.font.size = Pt(12)

# Numbered clauses
clauses = [
    'We (here in collectively referred as the "Partners") are carrying on the business as a partnership firm in the name and style of M/s IRIS COMMUNICATION having its registered office at (herein after referred to as the "firm") are desirous of availing credit facility in the form of credit facilities aggregating to INR 2000000 (Rupees Twenty Lakhs only) (the "facility") for business purposes from Kisetsu Saison Finance (India) Private Limited and Progfin Private Limited (formerly known as "Hytone Holdings Private Limited") (herein after together referred to as "Lenders").',
    "We shall advise in writing of any change that may take place in the said partnership firm and all the present / future partners shall be liable to you on any obligation which may be standing in the partnership firm's name in your books on the date of the receipt of such notice and until all such obligations shall have been liquidated.",
    "Mrs HARINI SANTOSH & Mrs NEETHI JAIN M S of the firm, be and are hereby severally authorized to: Convey acceptance of the sanction letter(s) to Lenders.",
    "Negotiation, finalise, settle, approve and convey the terms of the facility, in a form and manner acceptable to Lenders.",
    "Execute the finance and security documents with respect to the facility (including but not limited to the application(s), sanction letter(s), credit facility agreement, declaration, deed of hypothecation, deed of pledge, power(s) of attorney, as may be applicable) or any amendments or supplements there to as may be required by Lenders from time to time;",
    "Request such person(s) indicate in the sanction letter(s) indicated in the sanction letter(s) to give guarantees and security, if applicable;",
    "Seek extension / renewals of the facility on such terms and conditions as may be requested by the Lenders.",
    "Execute / furnish / authenticate / certify / collect / acknowledge / submit all deeds, documents, undertakings, declarations letters, application, statements, outstanding dues, acknowledgement, of debt in respect of the facility as may be required by Lenders from time to time; and",
    "Execute / complete all statutory, regulatory, and other forms and documents for availing / extending / renewing the facility and providing the security, if any, in the manner acceptable to the Lenders.",
    "We hereby certify that the firm is a registered firm bearing registration no 29AAIFI7005B1ZN under the provision of the Indian partnership Act 1932 and complying all the legal requirements of the Act.",
    "This authority shall continue to be force until all of us revoke it by a notice in writing delivered to the Lenders.",
]

for clause in clauses:
    doc.add_paragraph(clause, style="List Number")

# Share holding heading
sh = doc.add_paragraph()
shrun = sh.add_run("Share holding pattern and list of partners as on 20/05/2026 are as follows;")
shrun.bold = True

# Table
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "S.No."
hdr[1].text = "Name of Partner"
hdr[2].text = "Shareholding"
hdr[3].text = "Address"
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

rows_data = [
    ("1", "HARINI SANTOSH", "50%", "No.583, 4th A Cross, 4th Main, Marithilayout, Vasanthapura, Bengaluru - 560061"),
    ("2", "NEETHI JAIN M S", "50%", "No.M315, 5th Main, 3rd Cross, H Block, Ramakrishna Nagar, Mysore - 570023"),
]
for r in rows_data:
    cells = table.add_row().cells
    for i, val in enumerate(r):
        cells[i].text = val

doc.add_paragraph()

# Signature block as a 2-column table
sig = doc.add_table(rows=1, cols=2)
sig.style = "Table Grid"
c = sig.rows[0].cells
c[0].text = (
    "1. Signature of Partner\n"
    "Name: HARINI SANTOSH\n"
    "Date: 20/05/26\n"
    "For IRIS COMMUNICATION\n"
    "(signed)\n"
    "PARTNER"
)
c[1].text = (
    "2. Signature of Partner\n"
    "Name: Neethi Jain M S\n"
    "Date: 20/05/26\n"
    "For IRIS COMMUNICATION\n"
    "(signed)\n"
    "PARTNER"
)

doc.save("Friends loyal Document.docx")
print("Saved Friends loyal Document.docx")
